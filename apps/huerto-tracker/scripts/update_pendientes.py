"""Actualiza huerto-tracker-pendientes-v2.docx con el estado de la sesion 19/05/2026."""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os

SRC = "C:/WorkSpace/portfolio-apps/huerto-tracker-pendientes-v2.docx"
DST = "C:/WorkSpace/portfolio-apps/huerto-tracker-pendientes-v2.docx"

doc = docx.Document(SRC)

# -- helpers ------------------------------------------------------------------

def set_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def bullet(doc, text, done=False, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    if done:
        run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
    elif color:
        run.font.color.rgb = color
    return p

def normal(doc, text):
    return doc.add_paragraph(text, style="Normal")

# -- build new document -------------------------------------------------------

new = docx.Document()

# Copy styles from original (fonts, margins)
for style in doc.styles:
    try:
        new.styles[style.name]
    except:
        pass

# Reuse section margins from original
new._body.clear_content()

# 1. Title
p = new.add_heading("Huerto Tracker — Pendientes y notas importantes", level=1)

normal(new, "Actualizado: 19/05/2026 (Sprint 8)  ·  Branch: huerto  ·  Commits hasta 168e926")

# 2. Bugs conocidos
new.add_heading("1. Bugs conocidos", level=1)
new.add_heading("1.1 RESUELTOS", level=2)

new.add_heading("P2 — Medio", level=3)
bullet(new, "[HECHO] QuickLogModal.tsx — sustituido useCollection('diary_entries') por createStore; ya no carga todas las entradas en memoria.", done=True)
bullet(new, "[HECHO] diary.tsx — filtro gardenId activo en entradas, plantsById y exportacion CSV.", done=True)
bullet(new, "[HECHO] index.tsx (tabs) — filtros gardenId en yearHarvestKg, activeReminders y nextReminder.", done=True)

new.add_heading("P3 — Menor", level=3)
bullet(new, "[HECHO] plant/[id].tsx — batchNum corregido: cuenta plantas hermanas (misma cropId + gardenId), no entradas.", done=True)
bullet(new, "[HECHO] settings.tsx — URL App Store extraida a constante APP_STORE_URL.", done=True)

new.add_heading("1.2 Pendientes (bugs menores)", level=2)
new.add_heading("P3 — Menor", level=3)
bullet(new, "[P3] onboarding.tsx — LATAM_COUNTRIES con nombres en español (nombres propios, aceptable por ahora).")

# 3. Features implementadas
new.add_heading("2. Features implementadas en sesiones anteriores", level=1)

new.add_heading("Sprint 6 (completo)", level=2)
sprint6 = [
    "Riego inteligente (getNeedsWater, QuickLogModal)",
    "Cosecha rápida (harvesting flow)",
    "Compartir mapa del huerto (ViewShot + expo-sharing)",
    "Fotos de planta (photoUri)",
    "Guía offline de enfermedades (disease-guide con imágenes)",
    "Alertas de heladas (frost alert en home)",
    "Múltiples huertos (multi-garden con useActiveGarden)",
]
for f in sprint6:
    bullet(new, f"[HECHO] {f}", done=True)

new.add_heading("Sprint 8 — 19/05/2026 — Features GrowIt-inspired", level=2)
sprint8 = [
    "Onboarding wizard 7 pasos (welcome → space → method → sunlight → experience → provincia → huerto → éxito). Modelo UserProfile + hook useUserProfile singleton. i18n x6 locales.",
    "Plant detail tabs: Datos | Calendario | Vecinos | Cómo cultivar. DifficultyGauge 1-5 derivada de water/sun/spacing. CalendarGantt mes-a-mes con sowing/harvest por climateZone. HowToStages 5 etapas.",
    "Companion overlay en mapa: marcadores ✓ verdes (companion) / ✗ rojos (incompatible) en los bordes adyacentes (right/bottom) usando getCompatibilityStatus.",
    "Trial reminder switch en paywall: toggle 'avisarme antes de fin del trial' (default on). Tras compra programa notif local 6 días después (1 día antes de fin trial 7d) via scheduleDateAlert. AsyncStorage para preferencia.",
    "Month chips en plant cards: '🌱 Mar, Abr, May' debajo de variedad para plantas sin sowingDate. Lee crop.sowingMonths según climateZone activo.",
]
for f in sprint8:
    bullet(new, f"[HECHO] {f}", done=True)

new.add_heading("Sesión 19/05/2026 (mañana) — Funcionalidades nuevas", level=2)
session_done = [
    "Cultivos personalizados — modelo CustomCrop, hook useCustomCrops, screens /crop/index + /crop/new, botón en Settings. Integrados en plant/new.tsx y plant/[id].tsx.",
    "Drag & drop en mapa del huerto — react-native-reanimated v3.17.5 + GestureHandlerRootView. Long-press 280ms activa Pan gesture. Ghost animado sigue el dedo (UI thread). Drop: swapCells o setCell. scrollEnabled=false durante drag.",
    "Badge Nx en celdas del mapa — cada celda muestra cuántas veces aparece ese cultivo en el grid (ej: '9x', '4x'), inspirado en Planter.",
    "Color personalizado del huerto — 8 swatches de colores tierra en garden/edit.tsx. Campo color? en modelo Garden. Color visible en cabecera del grid y en imagen compartida.",
    "FAB 'Añadir planta' — botón flotante abajo-derecha en mapa. Auto-selecciona primera celda vacía.",
    "Notas del huerto — icono 📄 en header del mapa, bottom sheet con TextInput. Campo notes? en modelo Garden, persiste en AsyncStorage.",
    "Share mejorado para Instagram — ViewShot captura cabecera con color del huerto + nombre + stats + footer '🌱 Huerto Tracker'.",
]
for f in session_done:
    bullet(new, f"[HECHO] {f}", done=True)

# 4. Features pendientes
new.add_heading("3. Features pendientes / backlog", level=1)
new.add_heading("3.1 Alta prioridad", level=2)
bullet(new, "[Alta] Galería de fotos por planta — actualmente solo un photoUri. Historial de fotos mostraría evolución visual.")
bullet(new, "[Alta] Seed wishlist — lista de deseos de semillas/cultivos a comprar.")

new.add_heading("3.2 Media prioridad", level=2)
bullet(new, "[Media] syncAll.ts — existe pero no activo en producción. Revisar endpoints, conflictos y tests antes de activar.")
bullet(new, "[Test] CSV export del diario — verificar en Android (compartir fichero).")
bullet(new, "[Test] PDF informe estacional (usePdfReport) — testear con datos reales en dispositivo.")

new.add_heading("3.3 Mejoras mapa (backlog, inspiradas en Planter)", level=2)
bullet(new, "[Backlog] Ajuste dimensiones grid inline — modal con +/- en cada lado para añadir/quitar filas/cols desde cada borde.")
bullet(new, "[Backlog] Undo/Redo en mapa — deshacer últimos movimientos de plantas.")
bullet(new, "[Backlog] Garden Summary screen — resumen: tamaño, plantas por cultivo con cantidades totales.")

new.add_heading("3.4 Sprint 9 propuesto (GrowIt roadmap)", level=2)
bullet(new, "[Sprint 9] Plant scan AI con auto-fill: cámara → IA → rellena planting date, watering, place, growth stage automáticamente.")
bullet(new, "[Sprint 9] FAQ por crop — banco 40 preguntas frecuentes por cultivo (i18n key cropFaq.{id}.[]).")
bullet(new, "[Sprint 9] Recetas IA post-cosecha — al registrar harvest, sugerir recetas con ese ingrediente.")
bullet(new, "[Sprint 9] Tab Explore con categorías temáticas — discovery feed (hierbas medicinales, salsas, ensaladas...).")
bullet(new, "[Sprint 9] Snap tips contextual antes de fotografiar plantas (overlay con tips de framing/luz).")

new.add_heading("3.5 Sprint 10-11 propuesto", level=2)
bullet(new, "[Sprint 10] Reminders agrupados por bucket temporal (In 2 days / In 14 days) y tipo (Water/Progress).")
bullet(new, "[Sprint 10] Smart reminders auto basados en growth stage + clima.")
bullet(new, "[Sprint 10] Like/dislike feedback en sugerencias.")
bullet(new, "[Sprint 10] Tutorial mapa modal primera vez.")
bullet(new, "[Sprint 10] Storage/propagation post-harvest (almacenamiento, esquejes).")
bullet(new, "[Sprint 10] Clear cache con tamaño en Settings.")
bullet(new, "[Sprint 11] FAB central en tab bar + pulido visual.")

# 5. Antes de publicar
new.add_heading("4. Antes de publicar en App Store", level=1)
new.add_heading("4.1 Git / CI", level=2)
bullet(new, "[HECHO] git push origin huerto — hecho (16/05/2026)", done=True)
bullet(new, "Abrir PR: github.com/MrZamudioRamos/portfolio-apps/compare/huerto")
bullet(new, "[Seguridad] Revisar que no haya secrets en el diff antes de mergear.")

new.add_heading("4.2 Config de producción", level=2)
bullet(new, "[BLOQUEANTE] Reemplazar APP_STORE_URL en settings.tsx: https://apps.apple.com/app/id<APP_STORE_ID>")
bullet(new, "[Config] EXPO_PUBLIC_DEV_PRO=false en build de producción.")
bullet(new, "[Config] Claves RevenueCat para producción en .env.")
bullet(new, "[Config] API key para identifyPest (AI diagnosis de plagas).")
bullet(new, "[Verificar] app.json: version, bundleIdentifier, icono, splash.")

new.add_heading("4.3 QA", level=2)
bullet(new, "[QA] Flujo completo de compra en TestFlight / sandbox App Store Connect.")
bullet(new, "[QA] Drag & drop en dispositivo físico (gestos nativos difieren del simulador).")
bullet(new, "[QA] Cultivos personalizados: crear, editar, eliminar, usar en nueva planta.")
bullet(new, "[QA] Notas del huerto: guardar, editar, visible tras reinicio.")
bullet(new, "[QA] Share: imagen generada correcta en iOS y Android.")

# 6. Stack
new.add_heading("5. Stack (referencia)", level=1)
stack = [
    "Expo SDK 55  ·  expo-router file-based  ·  rama: huerto",
    "@portfolio/storage / notifications / billing / supabase / ui / shared",
    "Billing: Monthly 2,99 EUR / Annual 19,99 EUR",
    "expo-glass-effect (iOS 26+ GlassView, fallback View)",
    "react-native-gesture-handler ~2.30.0  ·  react-native-reanimated ~3.17.5",
    "i18n: 6 locales es/en/ca/gl/eu/val — completos",
    "77/77 vitest tests",
]
for s in stack:
    bullet(new, s)

# 7. Historial de sesiones
new.add_heading("6. Historial de sesiones", level=1)
sessions = [
    ("19/05/2026 (Sprint 8)", "Onboarding wizard 7 pasos, plant detail tabs, companion overlay en mapa, trial reminder switch, month chips en plant cards"),
    ("19/05/2026 (Sprint 7)", "Drag & drop mapa, badge Nx, color huerto, FAB, notas, share mejorado, cultivos personalizados (P2+P3 fix)"),
    ("16/05/2026", "Sprint 6 cerrado: riego, cosecha, fotos, enferm., heladas, multi-huerto, i18n P2, map isolation"),
    ("15/05/2026", "Mapa huerto v1, rotation, companions, stats, backup, notificaciones"),
]
for date, desc in sessions:
    bullet(new, f"{date} — {desc}")

new.save(DST)
print(f"Guardado: {DST}")
